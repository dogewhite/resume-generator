from docxtpl import DocxTemplate
import json
import os
from docx import Document

def check_template_content(template_path):
    """检查Word模板中的内容"""
    print("\n📝 检查模板内容...")
    try:
        doc = Document(template_path)
        print("\n模板中的文本内容:")
        for para in doc.paragraphs:
            if "{{" in para.text or "{%" in para.text:
                print(f"段落: {para.text}")
        
        print("\n表格中的内容:")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{" in cell.text or "{%" in cell.text:
                        print(f"单元格: {cell.text}")
    except Exception as e:
        print(f"❌ 读取模板内容时发生错误: {e}")

def test_template():
    # 检查文件是否存在
    if not os.path.exists("template1.docx"):
        print("❌ 错误: template1.docx 不存在")
        return
        
    if not os.path.exists("测试.json"):
        print("❌ 错误: 测试.json 不存在")
        return
    
    try:
        # 读取JSON数据
        print("📝 正在读取JSON数据...")
        with open("测试.json", "r", encoding="utf-8") as f:
            data = json.load(f)
        print("✅ JSON数据读取成功")
        
        # 检查模板内容
        check_template_content("template1.docx")
        
        # 加载模板
        print("\n📝 正在加载模板...")
        doc = DocxTemplate("template1.docx")
        
        # 获取模板中的变量
        print("\n📝 模板中的变量:")
        try:
            variables = doc.get_undeclared_template_variables()
            print(json.dumps(list(variables), indent=2, ensure_ascii=False))
        except Exception as e:
            print(f"❌ 获取模板变量时发生错误: {e}")
            print("💡 提示: 这可能是由于模板中的Jinja2语法错误导致")
            print("常见的错误包括:")
            print("1. 未闭合的循环 ({% for ... %} 没有对应的 {% endfor %})")
            print("2. 未闭合的条件语句 ({% if ... %} 没有对应的 {% endif %})")
            print("3. 变量名称错误")
            print("4. 语法错误 (例如使用了未定义的标签或过滤器)")
            return
        
        # 检查JSON数据中是否包含所有需要的变量
        print("\n📝 检查数据完整性...")
        missing_vars = []
        for var in variables:
            parts = var.split('.')
            value = data
            try:
                for part in parts:
                    value = value[part]
            except (KeyError, TypeError):
                missing_vars.append(var)
        
        if missing_vars:
            print("⚠️ 警告: 以下变量在JSON中缺失:")
            for var in missing_vars:
                print(f"  - {var}")
        else:
            print("✅ 所有必需的变量都存在")
        
        # 显示将要使用的数据
        print("\n📝 将要使用的数据:")
        print(json.dumps(data, indent=2, ensure_ascii=False))
        
        # 生成文档
        print("\n📝 正在生成文档...")
        doc.render(data)
        
        # 保存文档
        output_file = "测试报告.docx"
        doc.save(output_file)
        print(f"✅ 文档生成成功: {output_file}")
        
    except json.JSONDecodeError as e:
        print(f"❌ JSON解析错误: {e}")
    except Exception as e:
        print(f"❌ 发生错误: {e}")
        print("\n💡 错误分析:")
        if "Unexpected end of template" in str(e):
            print("模板中可能有未闭合的标签，请检查:")
            print("1. {% for ... %} 是否都有对应的 {% endfor %}")
            print("2. {% if ... %} 是否都有对应的 {% endif %}")
        elif "Unknown tag" in str(e):
            print("模板中使用了未知的标签，支持的标签包括:")
            print("1. for 循环: {% for item in items %}")
            print("2. if 条件: {% if condition %}")
            print("3. 变量: {{ variable }}")

if __name__ == "__main__":
    test_template() 